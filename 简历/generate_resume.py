#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph(doc, text, bold=False, font_size=10.5):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    return p

def add_bullet_point(doc, text, font_size=10.5):
    """添加带项目符号的段落"""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(font_size)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def create_resume():
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading('李云娟 - 前端开发工程师', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('女 | 30岁 | 9年工作经验 | 本科 | CET-6 | 软件设计师资格证\n')
    run.font.size = Pt(10.5)
    run = info.add_run('电话: 153 2039 9604 | 邮箱: lyj_515@163.com')
    run.font.size = Pt(10.5)

    doc.add_paragraph('_' * 80)

    # 求职意向
    add_heading(doc, '求职意向', level=1)
    add_paragraph(doc, '职位：高级前端开发工程师 | 性质：全职 | 薪资：面议 | 到岗时间：随时')

    # 个人技能
    add_heading(doc, '个人技能', level=1)

    add_paragraph(doc, '前端框架', bold=True, font_size=11)
    add_bullet_point(doc, 'Vue 生态：熟练运用 Vue 2/3、Vuex、Vue Router，5+ 年实战经验')
    add_bullet_point(doc, 'React 生态：熟练运用 React 18、React Router、React Hooks，4+ 年实战经验')
    add_bullet_point(doc, '熟练运用 TypeScript、ES6+ 语法')

    add_paragraph(doc, '前端工程化', bold=True, font_size=11)
    add_bullet_point(doc, '熟练使用 Webpack、Vite 等构建工具')
    add_bullet_point(doc, '了解 Monorepo 架构（Yarn Workspaces）、微前端架构（micro-app、qiankun）')
    add_bullet_point(doc, '熟悉 Git 工作流、CICD 流程')

    add_paragraph(doc, '其他技术', bold=True, font_size=11)
    add_bullet_point(doc, '熟悉 WebSocket、WebRTC 实时通信技术')
    add_bullet_point(doc, '熟悉移动端开发（Flex/Rem 布局、Hybrid App、小程序）')
    add_bullet_point(doc, '了解 Node.js、Electron 桌面应用开发')

    # 工作经历
    add_heading(doc, '工作经历', level=1)

    add_paragraph(doc, 'Fintopia 集团（原北京高域海汇科技有限公司）', bold=True, font_size=11)
    add_paragraph(doc, '前端开发工程师 | 2020.4 - 2024.6', font_size=10)
    add_bullet_point(doc, '参与 YChat 客服 CRM 平台开发，服务于 6 个国家、200+ 坐席')
    add_bullet_point(doc, '负责实时通信系统（WebSocket + WebRTC）开发和优化')
    add_bullet_point(doc, '负责多国家架构设计和实现，代码复用率 60%')
    add_bullet_point(doc, '负责性能优化，首屏加载时间从 4.5s 优化到 1.5s（提升 67%）')

    add_paragraph(doc, '北京任买科技有限公司', bold=True, font_size=11)
    add_paragraph(doc, '高级开发工程师 | 2019.11 - 2020.3', font_size=10)
    add_bullet_point(doc, '负责 H5 项目和后台管理系统开发，重构 H5 项目实现一套代码多 APP 使用')

    add_paragraph(doc, '北京奇艺世纪科技有限公司（爱奇艺）', bold=True, font_size=11)
    add_paragraph(doc, '高级研发工程师 | 2019.4 - 2019.11', font_size=10)
    add_bullet_point(doc, '参与组件库开发，制定前端样式开发统一标准')

    add_paragraph(doc, '能力天空（北京）科技有限公司', bold=True, font_size=11)
    add_paragraph(doc, 'Web 前端工程师 | 2016.11 - 2019.3', font_size=10)
    add_bullet_point(doc, '负责公司迭代任务开发，从头学习小程序开发并完成任务')

    # 项目经验
    add_heading(doc, '项目经验', level=1)

    # YChat 项目
    add_paragraph(doc, 'YChat 客服 CRM 平台 ⭐⭐⭐⭐⭐', bold=True, font_size=12)
    add_paragraph(doc, '项目周期：2020.04 - 2024.06 | 技术栈：Vue 2.6、Vuex、TypeScript、WebSocket、WebRTC、Webpack 5', font_size=9)

    add_paragraph(doc, '项目简介：', bold=True, font_size=10.5)
    add_paragraph(doc, 'YChat 是 Fintopia 集团的客服 CRM 平台，为一线电话客服和在线客服提供统一工作台。项目覆盖国内、印尼、墨西哥等 6 个国家，日均服务 200+ 名坐席，处理 5000+ 次会话，月均工单量 15000+。')

    add_paragraph(doc, '技术架构：', bold=True, font_size=10.5)
    add_bullet_point(doc, 'Monorepo 架构：Yarn Workspaces 管理 2300+ 文件，代码复用率 60%')
    add_bullet_point(doc, '微前端架构：micro-app 集成子应用，JS/CSS 自动隔离')
    add_bullet_point(doc, '多国家架构：路由/配置/灰度三层控制，一套代码支持 6 个国家')
    add_bullet_point(doc, '实时通信：WebSocket (IM) + WebRTC (电话) 双通道')

    add_paragraph(doc, '我负责的核心模块：', bold=True, font_size=10.5)

    add_paragraph(doc, '1. 实时通信系统', bold=True)
    add_bullet_point(doc, '实现指数退避重连策略和消息队列缓存机制')
    add_bullet_point(doc, '重连成功率从 80% 提升到 98%，消息丢失率从 5% 降低到 0.1%')
    add_bullet_point(doc, '使用 JsSIP + SIP 协议实现 WebRTC 电话通信')
    add_bullet_point(doc, '使用 Broadcast Channel 解决跨 Tab 冲突')

    add_paragraph(doc, '2. 多国家架构', bold=True)
    add_bullet_point(doc, '路由分叉：路由 Meta 声明 businessCode，守卫自动拦截')
    add_bullet_point(doc, '配置中心：运行时动态加载国家配置，支持热更新')
    add_bullet_point(doc, '国际化：支持 5 种语言（中英印尼西葡）')

    add_paragraph(doc, '3. 新工作台', bold=True)
    add_bullet_point(doc, '三栏布局：客户信息 + 对话区 + 工单面板，一屏展示所有信息')
    add_bullet_point(doc, 'bizSource 驱动：通过参数切换在线/电话模式')
    add_bullet_point(doc, '坐席效率提升 30%，工单创建时间减少 40%，用户满意度提升 15%')

    add_paragraph(doc, '4. 性能优化', bold=True)
    add_bullet_point(doc, '首屏加载：4.5s → 1.5s（提升 67%），Bundle：3.2MB → 1.1MB（减少 65%）')
    add_bullet_point(doc, '长列表优化：DOM 节点从 1000+ 降低到 20（减少 98%）')
    add_bullet_point(doc, '使用虚拟滚动、懒加载、代码分割等技术')

    add_paragraph(doc, '项目成果：', bold=True, font_size=10.5)
    add_bullet_point(doc, '支持 6 个国家业务，日均服务 200+ 坐席，系统可用性 99.9%')
    add_bullet_point(doc, '首屏加载优化 67%，Bundle 体积优化 65%')
    add_bullet_point(doc, 'WebSocket 重连成功率 98%，消息丢失率 0.1%')
    add_bullet_point(doc, '新工作台坐席效率提升 30%')

    # 其他项目（简化）
    add_paragraph(doc, '叨叨桌面端聊天工具', bold=True, font_size=11)
    add_paragraph(doc, '项目周期：2020.04 - 2024.06 | 技术栈：React、Vue、Electron', font_size=9)
    add_paragraph(doc, '面向公司内部使用的桌面端聊天工具，负责编辑器开发、聊天窗口渲染、多选/预览功能、性能优化等。')

    add_paragraph(doc, 'IronBank PC 端项目', bold=True, font_size=11)
    add_paragraph(doc, '项目周期：2022.06 - 2024.06 | 技术栈：React、React Hooks', font_size=9)
    add_paragraph(doc, '定制的后台管理系统，用于公司内部项目管理和周报管理。')

    add_paragraph(doc, '任意花 H5（已上线）', bold=True, font_size=11)
    add_paragraph(doc, '项目周期：2019.11 - 2020.03 | 技术栈：Vue、Nuxt.js', font_size=9)
    add_paragraph(doc, '贷款完整流程的 H5 项目，嵌入 APP 使用，包括实名认证、授信、提现、还款等流程。')

    add_paragraph(doc, '爱奇艺号 PC 端（已上线）', bold=True, font_size=11)
    add_paragraph(doc, '项目周期：2019.4 - 2019.11 | 技术栈：Vue、Vuex', font_size=9)
    add_paragraph(doc, '爱奇艺旗下视频内容创作、分发、变现平台。负责自媒体模块开发、组件库开发、样式规范制定。')

    add_paragraph(doc, '能力天空相关项目（PC 端、小程序、APP）', bold=True, font_size=11)
    add_paragraph(doc, '项目周期：2016.11 - 2019.3 | 技术栈：Vue、jQuery、小程序', font_size=9)
    add_paragraph(doc, '在线教育平台，负责 PC 端、小程序、APP 的开发和维护。')

    # 教育经历
    add_heading(doc, '教育经历', level=1)
    add_paragraph(doc, '东北石油大学 | 软件工程（统招·一本）| 2013.9 - 2017.6')

    # 自我评价
    add_heading(doc, '自我评价', level=1)
    add_bullet_point(doc, '9 年前端开发经验，具备扎实的前端基础和丰富的项目经验')
    add_bullet_point(doc, '熟练掌握 Vue、React 生态，有大型项目（2300+ 文件）的架构设计和优化经验')
    add_bullet_point(doc, '具备 Monorepo、微前端、多国家架构等复杂架构实战经验')
    add_bullet_point(doc, '擅长性能优化，有首屏加载、长列表、Bundle 体积等多方面的优化经验')
    add_bullet_point(doc, '熟悉 WebSocket、WebRTC 实时通信技术，有成熟的解决方案')
    add_bullet_point(doc, '学习能力强，喜欢研究新技术，团队协作能力强，工作态度认真负责')

    # 保存文档
    doc.save('/Users/lww/ayy/简历/李云娟的前端简历-2026新版.doc')
    print('简历生成成功：李云娟的前端简历-2026新版.doc')

if __name__ == '__main__':
    create_resume()
